import concurrent.futures
import io
import os
import pathlib
import stat
import tarfile
import tempfile
import unicodedata
import zipfile
from datetime import datetime

import fitz
import py7zr
from docx import Document
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, ImageOps

from app.core.config import settings


class ProcessingLimitError(ValueError):
    """The submitted document would exceed a server-side processing limit."""


class UnsafeArchiveError(ValueError):
    """The submitted archive contains an unsafe or unsupported member."""


class ArchiveLimitError(ProcessingLimitError):
    """The expanded archive would exceed a server-side processing limit."""


class _ArchiveBudget:
    def __init__(self):
        self.count = 0
        self.total_size = 0
        self.names: set[str] = set()

    def add(self, raw_name: str, size: int) -> str:
        name = _normalize_archive_name(raw_name)
        folded = name.casefold()
        if folded in self.names:
            raise UnsafeArchiveError("Archive contains duplicate paths")
        if size < 0:
            raise UnsafeArchiveError("Archive contains an invalid file size")

        self.count += 1
        self.total_size += size
        if self.count > settings.WATERMARK_MAX_ARCHIVE_FILES:
            raise ArchiveLimitError("Archive contains too many files")
        if size > settings.WATERMARK_MAX_ARCHIVE_FILE_SIZE:
            raise ArchiveLimitError("Archive member is too large")
        if self.total_size > settings.WATERMARK_MAX_ARCHIVE_TOTAL_SIZE:
            raise ArchiveLimitError("Archive expands beyond the allowed total size")

        self.names.add(folded)
        return name


def _normalize_archive_name(raw_name: str) -> str:
    value = str(raw_name or "").replace("\\", "/")
    if not value or value.startswith("/") or "\x00" in value:
        raise UnsafeArchiveError("Archive contains an invalid path")
    if any(unicodedata.category(char).startswith("C") for char in value):
        raise UnsafeArchiveError("Archive paths may not contain control characters")

    path = pathlib.PurePosixPath(value)
    parts = path.parts
    if (
        path.is_absolute()
        or not parts
        or any(part in {"", ".", ".."} for part in parts)
        or parts[0].endswith(":")
        or len(parts) > settings.MEDIA_MAX_DEPTH
        or any(len(part) > settings.ADMIN_MAX_FILENAME_LENGTH for part in parts)
    ):
        raise UnsafeArchiveError("Archive contains a path outside its root")
    return path.as_posix()


def _read_limited(stream, expected_size: int) -> bytes:
    output = io.BytesIO()
    remaining = min(expected_size, settings.WATERMARK_MAX_ARCHIVE_FILE_SIZE) + 1
    while remaining > 0:
        chunk = stream.read(min(1024 * 1024, remaining))
        if not chunk:
            break
        output.write(chunk)
        remaining -= len(chunk)
    data = output.getvalue()
    if len(data) != expected_size:
        raise UnsafeArchiveError("Archive member size does not match its metadata")
    return data


def _reject_suspicious_ratio(uncompressed: int, compressed: int | None) -> None:
    if compressed is None:
        return
    compressed_size = int(compressed)
    if uncompressed > 1024 * 1024 and uncompressed > max(compressed_size, 1) * 100:
        raise ArchiveLimitError("Archive compression ratio is too high")


def _zip_members(archive_io: io.BytesIO) -> list[tuple[str, bytes]]:
    result = []
    budget = _ArchiveBudget()
    with zipfile.ZipFile(archive_io) as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            if info.flag_bits & 0x1:
                raise UnsafeArchiveError("Encrypted archives are not supported")
            mode = info.external_attr >> 16
            if stat.S_ISLNK(mode):
                raise UnsafeArchiveError("Archive links are not allowed")
            name = budget.add(info.filename, info.file_size)
            _reject_suspicious_ratio(info.file_size, info.compress_size)
            with archive.open(info, "r") as member:
                result.append((name, _read_limited(member, info.file_size)))
    return result


def _tar_members(archive_io: io.BytesIO) -> list[tuple[str, bytes]]:
    result = []
    budget = _ArchiveBudget()
    with tarfile.open(fileobj=archive_io, mode="r:*") as archive:
        for member in archive:
            if member.isdir():
                continue
            if not member.isfile():
                raise UnsafeArchiveError("Archive links and special files are not allowed")
            name = budget.add(member.name, member.size)
            source = archive.extractfile(member)
            if source is None:
                raise UnsafeArchiveError("Archive member could not be read")
            with source:
                result.append((name, _read_limited(source, member.size)))
    return result


def _seven_zip_members(archive_io: io.BytesIO) -> list[tuple[str, bytes]]:
    budget = _ArchiveBudget()
    expected: dict[str, int] = {}
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = pathlib.Path(tmp_dir).resolve()
        with py7zr.SevenZipFile(archive_io, mode="r") as archive:
            for info in archive.list():
                if info.is_directory:
                    continue
                if getattr(info, "is_symlink", False) or not getattr(info, "is_file", True):
                    raise UnsafeArchiveError("Archive links and special files are not allowed")
                name = budget.add(info.filename, int(info.uncompressed))
                _reject_suspicious_ratio(int(info.uncompressed), getattr(info, "compressed", None))
                expected[name] = int(info.uncompressed)
            if expected:
                archive.extract(path=tmp_path, targets=list(expected))

        result = []
        found = set()
        for path in tmp_path.rglob("*"):
            if path.is_symlink():
                raise UnsafeArchiveError("Archive links are not allowed")
            if not path.is_file():
                continue
            resolved = path.resolve(strict=True)
            if not resolved.is_relative_to(tmp_path):
                raise UnsafeArchiveError("Archive extracted outside its root")
            name = path.relative_to(tmp_path).as_posix()
            normalized = _normalize_archive_name(name)
            if normalized not in expected or normalized in found:
                raise UnsafeArchiveError("Archive extraction produced unexpected files")
            with path.open("rb") as source:
                result.append((normalized, _read_limited(source, expected[normalized])))
            found.add(normalized)
        if found != set(expected):
            raise UnsafeArchiveError("Archive is incomplete or invalid")
        return result


def _validate_docx_container(content: bytes) -> None:
    budget = _ArchiveBudget()
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            for info in archive.infolist():
                if info.is_dir():
                    continue
                mode = info.external_attr >> 16
                if info.flag_bits & 0x1 or stat.S_ISLNK(mode):
                    raise UnsafeArchiveError("Document package contains unsupported entries")
                budget.add(info.filename, info.file_size)
                _reject_suspicious_ratio(info.file_size, info.compress_size)
    except zipfile.BadZipFile as exc:
        raise UnsafeArchiveError("Invalid Word document package") from exc


def get_font_path():
    # Prefer the path supplied by the container environment.
    env_path = os.getenv("WATERMARK_FONT_PATH")
    if env_path and pathlib.Path(env_path).exists():
        return env_path

    # Fall back to the standard Windows font location for local execution.
    win_path = "C:\\Windows\\Fonts\\simhei.ttf"
    if pathlib.Path(win_path).exists():
        return win_path

    # Pillow falls back to its default font when neither path is available.
    return None


# --- 1. Image processor: full-frame tiling and EXIF orientation ---
def process_single_image(content: bytes, text: str) -> bytes:
    try:
        # Normalize EXIF orientation before calculating the output dimensions.
        raw_img = Image.open(io.BytesIO(content))
        raw_width, raw_height = raw_img.size
        if (
            raw_width <= 0
            or raw_height <= 0
            or raw_width * raw_height > settings.WATERMARK_MAX_IMAGE_PIXELS
        ):
            raise ProcessingLimitError("Image dimensions exceed the processing limit")
        img = ImageOps.exif_transpose(raw_img).convert("RGBA")
        width, height = img.size

        # Include a minute-level timestamp in the watermark text.
        full_text = f"{text} {datetime.now().strftime('%Y-%m-%d %H:%M')}"

        # Scale the font against the longer image edge.
        base_side = max(width, height)
        font_size = int(base_side / 50)
        font_size = max(font_size, 15)

        try:
            font = ImageFont.truetype(get_font_path(), font_size)
        except Exception:
            font = ImageFont.load_default()

        # Measure the rendered text before selecting the tile spacing.
        draw_temp = ImageDraw.Draw(Image.new("RGBA", (1, 1)))
        bbox = draw_temp.textbbox((0, 0), full_text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]

        # Leave horizontal and vertical space between adjacent tiles.
        step_x = int(tw * 1.5)
        step_y = int(th * 4)

        # Use a double-sized canvas so rotation does not expose empty corners.
        overlay = Image.new("RGBA", (width * 2, height * 2), (255, 255, 255, 0))
        draw = ImageDraw.Draw(overlay)

        # Tile the text in semi-transparent gold.
        for x in range(0, width * 2, step_x):
            for y in range(0, height * 2, step_y):
                draw.text((x, y), full_text, fill=(255, 215, 0, 160), font=font)

        # Rotate by 45 degrees and crop back to the source dimensions.
        overlay = overlay.rotate(45, resample=Image.BICUBIC)
        left = (overlay.width - width) // 2
        top = (overlay.height - height) // 2
        txt_layer = overlay.crop((left, top, left + width, top + height))

        # Composite the watermark with the normalized source image.
        combined = Image.alpha_composite(img, txt_layer)
        out = io.BytesIO()
        combined.convert("RGB").save(out, format="JPEG", quality=90)
        return out.getvalue()
    except ProcessingLimitError:
        raise
    except Exception as e:
        print(f"Image processing failed: {e}")
        return content


def create_watermark_layer(width, height, text):
    """Create a transparent watermark layer matching a PDF page."""
    # Build the same 45-degree tiled watermark used for images.
    # full_text = f"{text} {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    full_text = text

    # Use a double-sized square canvas to cover the page after rotation.
    overlay_size = max(width, height) * 2
    overlay = Image.new("RGBA", (overlay_size, overlay_size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)

    # Scale the font and tile spacing to the canvas.
    font_size = max(int(overlay_size / 80), 15)
    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()

    # Draw the repeating watermark grid.
    for x in range(0, overlay_size, font_size * 10):
        for y in range(0, overlay_size, font_size * 5):
            draw.text((x, y), full_text, fill=(255, 215, 0, 100), font=font)

    # Rotate the grid before cropping it to the PDF page.
    overlay = overlay.rotate(45, resample=Image.BICUBIC)

    # Select the centered region matching the target page.
    left = (overlay_size - width) // 2
    top = (overlay_size - height) // 2
    return overlay.crop((left, top, left + width, top + height))


def process_single_pdf(pdf_bytes: bytes, text: str) -> bytes:
    doc = None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count > settings.WATERMARK_MAX_PDF_PAGES:
            raise ProcessingLimitError("PDF contains too many pages")

        for page in doc:
            # PDF page dimensions are expressed in points.
            rect = page.rect
            w, h = int(rect.width), int(rect.height)
            if w <= 0 or h <= 0 or w * h > settings.WATERMARK_MAX_IMAGE_PIXELS:
                raise ProcessingLimitError("PDF page dimensions exceed the processing limit")

            # Generate the rotated watermark layer.
            watermark_img = create_watermark_layer(w, h, text)

            # Encode the layer as PNG bytes for PDF insertion.
            img_byte_arr = io.BytesIO()
            watermark_img.save(img_byte_arr, format="PNG")

            # Insert above existing page content.
            page.insert_image(rect, stream=img_byte_arr.getvalue(), overlay=True)

        return doc.write()
    except ProcessingLimitError:
        raise
    except Exception as e:
        print(f"PDF processing failed: {e}")
        return pdf_bytes
    finally:
        if doc is not None:
            doc.close()


# --- 3. Word processor ---
def process_single_word(content: bytes, text: str) -> bytes:
    try:
        _validate_docx_container(content)
        doc = Document(io.BytesIO(content))
        mark_text = f"{text} {datetime.now().strftime('%Y-%m-%d')}"
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = p.add_run(mark_text)
            run.font.size, run.font.color.rgb = Pt(24), RGBColor(255, 215, 0)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except (ProcessingLimitError, UnsafeArchiveError):
        raise
    except Exception:
        return content


# --- 4. Task dispatch and filtering ---
def dispatch_task(item, text):
    name, content = item
    path_obj = pathlib.Path(name)
    ext = path_obj.suffix.lower()

    # Ignore hidden metadata and operating-system artifacts.
    if name.startswith('__MACOSX') or path_obj.name.startswith('._') or ext == '.db':
        return name, content

    try:
        if ext in ['.jpg', '.jpeg', '.png']:
            # Ignore thumbnail-sized images smaller than 5 KiB.
            if len(content) < 5120: return name, content
            return name, process_single_image(content, text)
        elif ext == '.pdf':
            return name, process_single_pdf(content, text)
        elif ext in ['.docx', '.doc']:
            return name, process_single_word(content, text)
    except (ProcessingLimitError, UnsafeArchiveError):
        raise
    except Exception as e:
        print(f"Error in dispatcher for {name}: {e}")

    return name, content


def run_batch_task(files_data: list, text: str):
    """Process a batch concurrently while preserving input order."""
    if not files_data:
        return []
    workers = max(1, min(settings.WATERMARK_MAX_WORKERS, len(files_data)))
    with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
        # Executor.map preserves the same order as files_data.
        results = list(executor.map(lambda x: dispatch_task(x, text), files_data))
    return results


# --- 5. Generic archive processor ---
def process_any_archive(archive_bytes: bytes, text: str, archive_ext: str) -> bytes:
    archive_io = io.BytesIO(archive_bytes)
    try:
        if archive_ext == ".7z":
            files_to_work = _seven_zip_members(archive_io)
        elif archive_ext == ".zip":
            files_to_work = _zip_members(archive_io)
        elif archive_ext in {".tar", ".tar.gz", ".tar.bz2", ".tar.xz", ".tgz", ".tbz2", ".txz"}:
            files_to_work = _tar_members(archive_io)
        else:
            raise UnsafeArchiveError("Unsupported archive format")
    except (ProcessingLimitError, UnsafeArchiveError):
        raise
    except (zipfile.BadZipFile, tarfile.TarError, OSError, EOFError) as exc:
        raise UnsafeArchiveError("Invalid or unsupported archive") from exc
    except Exception as exc:
        # py7zr exposes format-specific exception types across versions. Keep
        # parser details out of the HTTP response and fail closed.
        raise UnsafeArchiveError("Invalid or unsupported archive") from exc

    processed_list = run_batch_task(files_to_work, text)
    out_buffer = io.BytesIO()
    with zipfile.ZipFile(out_buffer, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as out_zip:
        for name, data in processed_list:
            out_zip.writestr(name, data)
    return out_buffer.getvalue()
